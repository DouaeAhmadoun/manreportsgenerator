import streamlit as st
from datetime import datetime
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Any, Union


class WordUtils:
    """Classe contenant les utilitaires pour la manipulation de documents Word"""
    
    @staticmethod
    def format_success_rate(rate: Union[int, float]) -> str:
        """Formate le taux de réussite en pourcentage"""
        if isinstance(rate, (int, float)):
            return f"{rate:.1%}"
        return "0%"
    
    @staticmethod
    def format_date(date_str: Any) -> str:
        """Formate une date pour l'affichage"""
        if not date_str:
            return ""
        try:
            # Si c'est déjà une string de date formatée
            if isinstance(date_str, str) and len(date_str) == 10:
                return date_str
            # Sinon essayer de parser et reformater
            dt = datetime.fromisoformat(str(date_str))
            return dt.strftime("%d/%m/%Y")
        except:
            return str(date_str)
    
    @staticmethod
    def format_number(number: Union[int, float], decimals: int = 1) -> str:
        """Formate un nombre avec le bon nombre de décimales"""
        if isinstance(number, (int, float)):
            if decimals == 0:
                return f"{number:.0f}"
            else:
                return f"{number:.{decimals}f}"
        return "0"
    
    @staticmethod
    def format_percentage(value: Union[int, float], decimals: int = 1) -> str:
        """Formate une valeur en pourcentage"""
        if isinstance(value, (int, float)):
            return f"{value:.{decimals}f}%"
        return "0%"
    
    @staticmethod
    def add_landscape_section_for_annexes(doc):
        """Version corrigée : déplace le contenu des annexes dans une section paysage"""
        try:
            st.write("📄 **Recherche et déplacement des annexes...**")
            
            # 1. Trouver où commencent les annexes
            annexes_start_paragraph = None
            annexes_paragraphs = []
            found_annexes = False
            
            for paragraph in doc.paragraphs:
                paragraph_text = paragraph.text.strip().lower()
                
                # Détecter le début des annexes
                if not found_annexes and any(marker in paragraph_text for marker in [
                    "tableau des essais et commentaires",
                    "# tests",
                    "annexes trouvees",
                    "simulations trouvees"
                ]):
                    found_annexes = True
                    annexes_start_paragraph = paragraph
                
                # Collecter tous les paragraphes d'annexes
                if found_annexes:
                    annexes_paragraphs.append(paragraph)
            
            if annexes_start_paragraph and annexes_paragraphs:
                # 2. Insérer un saut de section AVANT les annexes
                WordUtils.insert_section_break_before_paragraph(doc, annexes_start_paragraph)
                
                # 3. Configurer la nouvelle section en paysage
                last_section = doc.sections[-1]
                last_section.orientation = WD_ORIENT.LANDSCAPE
                last_section.page_width = Inches(11.69)
                last_section.page_height = Inches(8.27)
                
                # Marges optimisées pour tableaux larges
                last_section.left_margin = Inches(0.4)
                last_section.right_margin = Inches(0.4)
                last_section.top_margin = Inches(0.3)
                last_section.bottom_margin = Inches(0.3)
                
                st.write(f"✅ **Section paysage créée avec {len(annexes_paragraphs)} paragraphes d'annexes**")
            else:
                st.warning("⚠️ **Annexes non détectées** - pas de changement d'orientation")
            
        except Exception as e:
            st.error(f"❌ **Erreur création section paysage :** {e}")
            import traceback
            st.error(traceback.format_exc())
    
    @staticmethod
    def insert_section_break_before_paragraph(doc, target_paragraph):
        """Insère un saut de section avant un paragraphe - Version robuste"""
        try:
            # Méthode 1: Essayer d'insérer un saut de section via le paragraphe
            try:
                # Trouver l'élément parent
                parent = target_paragraph._element.getparent()
                
                # Créer un nouveau paragraphe avec saut de section
                new_p = OxmlElement('w:p')
                
                # Ajouter les propriétés de paragraphe avec saut de section
                pPr = OxmlElement('w:pPr')
                sectPr = OxmlElement('w:sectPr')
                
                # Type de saut : nouvelle page
                sectPrChange = OxmlElement('w:type')
                sectPrChange.set(qn('w:val'), 'nextPage')
                sectPr.append(sectPrChange)
                
                pPr.append(sectPr)
                new_p.append(pPr)
                
                # Insérer avant le paragraphe cible
                parent.insert(parent.index(target_paragraph._element), new_p)
                return True
                
            except Exception as e1:
                st.warning(f"⚠️ **Méthode 1 échouée :** {e1}")
                
                # Méthode 2: Saut de page simple
                try:
                    # Créer un saut de page simple avant
                    page_break = OxmlElement('w:p')
                    page_break_run = OxmlElement('w:r')
                    page_break_elem = OxmlElement('w:br')
                    page_break_elem.set(qn('w:type'), 'page')
                    
                    page_break_run.append(page_break_elem)
                    page_break.append(page_break_run)
                    
                    # Insérer avant le paragraphe cible
                    parent = target_paragraph._element.getparent()
                    parent.insert(parent.index(target_paragraph._element), page_break)
                    
                    st.write("✅ **Saut de page inséré (méthode 2)**")
                    return True
                    
                except Exception as e2:
                    st.warning(f"⚠️ **Méthode 2 échouée :** {e2}")
                    return False
        
        except Exception as e:
            st.error(f"❌ **Erreur insertion saut :** {e}")
            return False
    
    @staticmethod
    def add_page_break(doc):
        """Ajoute un saut de page au document"""
        try:
            paragraph = doc.add_paragraph()
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
            return True
        except Exception as e:
            st.warning(f"Impossible d'ajouter un saut de page: {e}")
            return False
    
    @staticmethod
    def set_document_margins(doc, left=1.0, right=1.0, top=1.0, bottom=1.0):
        """Définit les marges du document en pouces"""
        try:
            for section in doc.sections:
                section.left_margin = Inches(left)
                section.right_margin = Inches(right)
                section.top_margin = Inches(top)
                section.bottom_margin = Inches(bottom)
            return True
        except Exception as e:
            st.warning(f"Impossible de définir les marges: {e}")
            return False
    
    @staticmethod
    def optimize_document_performance(doc):
        """Optimise les performances du document"""
        optimizations = []
        
        try:
            # Compter les éléments
            para_count = len(doc.paragraphs)
            table_count = len(doc.tables)
            
            if para_count > 1000:
                optimizations.append("Document très long - considérer la pagination")
            
            if table_count > 50:
                optimizations.append("Nombreux tableaux - vérifier les performances")
            
            return {
                "optimizations": optimizations,
                "performance_score": min(100, max(0, 100 - para_count // 10 - table_count * 2))
            }
            
        except Exception as e:
            return {"error": str(e), "optimizations": [], "performance_score": 50}
